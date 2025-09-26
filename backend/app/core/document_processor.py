"""Document processing utilities for extracting text from various file formats."""
from pathlib import Path
from typing import Dict, Any, Union
import logging

import pypdf
from docx import Document as DocxDocument
import pdfplumber

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Handles text extraction from various document formats."""

    SUPPORTED_FORMATS = {'.pdf', '.txt', '.docx', '.doc'}

    @staticmethod
    def extract_text(file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Extract text content from a document file.

        Args:
            file_path: Path to the document file

        Returns:
            Dict containing:
                - content: Extracted text content
                - format: File format
                - error: Error message if extraction failed
                - page_count: Number of pages (for PDFs)
                - word_count: Approximate word count
        """
        try:
            file_path = Path(file_path) if isinstance(file_path, str) else file_path
            if not file_path.exists():
                return {
                    "content": "",
                    "format": "",
                    "error": f"File not found: {file_path}",
                    "page_count": 0,
                    "word_count": 0
                }

            suffix = file_path.suffix.lower()

            if suffix == '.txt':
                return DocumentProcessor._extract_txt(file_path)
            elif suffix == '.pdf':
                return DocumentProcessor._extract_pdf(file_path)
            elif suffix in ['.docx', '.doc']:
                return DocumentProcessor._extract_docx(file_path)
            else:
                return {
                    "content": "",
                    "format": suffix,
                    "error": f"Unsupported file format: {suffix}",
                    "page_count": 0,
                    "word_count": 0
                }

        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return {
                "content": "",
                "format": file_path.suffix.lower() if file_path else "",  # pyright: ignore[reportAttributeAccessIssue]
                "error": str(e),
                "page_count": 0,
                "word_count": 0
            }

    @staticmethod
    def _extract_txt(file_path: Path) -> Dict[str, Any]:
        """Extract text from a plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()

            return {
                "content": content,
                "format": ".txt",
                "error": None,
                "page_count": 1,
                "word_count": len(content.split())
            }
        except Exception as e:
            raise Exception(f"Failed to read text file: {str(e)}")

    @staticmethod
    def _extract_pdf(file_path: Path) -> Dict[str, Any]:
        """Extract text from a PDF file using both pypdf and pdfplumber for better extraction."""
        content_parts = []
        page_count = 0

        # Try pdfplumber first (better for complex layouts)
        try:
            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        content_parts.append(text)
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed, trying pypdf: {str(e)}")

            # Fallback to pypdf
            try:
                with open(file_path, 'rb') as f:
                    pdf_reader = pypdf.PdfReader(f)
                    page_count = len(pdf_reader.pages)

                    for page in pdf_reader.pages:
                        text = page.extract_text()
                        if text:
                            content_parts.append(text)
            except Exception as e:
                raise Exception(f"Failed to extract PDF content: {str(e)}")

        content = "\n\n".join(content_parts)

        return {
            "content": content,
            "format": ".pdf",
            "error": None,
            "page_count": page_count,
            "word_count": len(content.split())
        }

    @staticmethod
    def _extract_docx(file_path: Path) -> Dict[str, Any]:
        """Extract text from a DOCX file."""
        try:
            doc = DocxDocument(str(file_path))
            content_parts = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content_parts.append(" | ".join(row_text))

            content = "\n\n".join(content_parts)

            return {
                "content": content,
                "format": ".docx",
                "error": None,
                "page_count": 1,  # DOCX doesn't have pages in the same way
                "word_count": len(content.split())
            }
        except Exception as e:
            raise Exception(f"Failed to extract DOCX content: {str(e)}")

    @staticmethod
    def prepare_for_llm(text: str, max_length: int = 10000) -> str:
        """
        Prepare extracted text for LLM consumption.

        Args:
            text: Raw extracted text
            max_length: Maximum character length (default 10000)

        Returns:
            Cleaned and truncated text suitable for LLM context
        """
        # Clean up excessive whitespace
        lines = text.split('\n')
        cleaned_lines = []

        for line in lines:
            line = line.strip()
            if line:  # Skip empty lines
                cleaned_lines.append(line)

        # Join with single newlines
        cleaned_text = '\n'.join(cleaned_lines)

        # Truncate if necessary
        if len(cleaned_text) > max_length:
            cleaned_text = cleaned_text[:max_length] + "...\n[Content truncated]"

        return cleaned_text

    @staticmethod
    def format_document_context(documents: list[Dict[str, Any]]) -> str:
        """
        Format multiple documents for inclusion in LLM context.

        Args:
            documents: List of document dictionaries with 'filename' and 'content' keys

        Returns:
            Formatted string for LLM context
        """
        if not documents:
            return ""

        context_parts = ["You have access to the following documents:\n"]

        for i, doc in enumerate(documents, 1):
            filename = doc.get('filename', f'Document {i}')
            content = doc.get('content', '')

            context_parts.append(f"\n--- Document {i}: {filename} ---\n")
            context_parts.append(DocumentProcessor.prepare_for_llm(content))
            context_parts.append("\n--- End of Document ---\n")

        return ''.join(context_parts)
